from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
from enum import StrEnum
from io import BytesIO
from pathlib import Path

from app.services.ingestion.filenames import safe_filename

MIN_EXTRACTED_TEXT_CHARS = 50

PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
TXT_MIME_TYPES = {"text/plain", "text/markdown", "application/octet-stream"}


class IngestionError(ValueError):
    """Raised when source content cannot be safely ingested."""


class SourceKind(StrEnum):
    PDF = "pdf"
    DOCX = "docx"
    TEXT = "text"
    URL = "url"


@dataclass(slots=True)
class ExtractedText:
    """Text extracted from an uploaded file or remote source."""

    text: str
    source_kind: SourceKind
    filename: str | None = None
    content_type: str | None = None
    warnings: list[str] = field(default_factory=list)

    @property
    def char_count(self) -> int:
        return len(self.text)


def extract_uploaded_text(
    file_bytes: bytes,
    *,
    filename: str | None,
    content_type: str | None,
) -> ExtractedText:
    """Extract plain text from a PDF, DOCX, or text upload."""

    if not file_bytes:
        raise IngestionError("Uploaded file is empty.")

    extension = Path(filename or "").suffix.lower()
    normalized_type = (content_type or "").split(";", maxsplit=1)[0].strip().lower()

    if extension == ".pdf" or normalized_type in PDF_MIME_TYPES:
        return extract_pdf_text(file_bytes, filename=filename, content_type=content_type)
    if extension == ".docx" or normalized_type in DOCX_MIME_TYPES:
        return extract_docx_text(file_bytes, filename=filename, content_type=content_type)
    if extension in {".txt", ".md", ".text"} or normalized_type in TXT_MIME_TYPES:
        return extract_txt_text(file_bytes, filename=filename, content_type=content_type)

    raise IngestionError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")


def extract_pdf_text(
    file_bytes: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> ExtractedText:
    """Extract text from a PDF file using pypdf."""

    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise IngestionError("PDF extraction requires the pypdf package.") from exc

    try:
        reader = PdfReader(BytesIO(file_bytes))
        page_texts = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # pragma: no cover - library-specific failures
        raise IngestionError("Could not read the uploaded PDF file.") from exc

    warnings: list[str] = []
    if not page_texts:
        warnings.append("PDF has no readable pages.")
    empty_pages = sum(1 for text in page_texts if not text.strip())
    if empty_pages:
        warnings.append(
            f"{empty_pages} PDF page(s) had no extractable text; scanned PDFs may need OCR."
        )

    return _validated_extraction(
        "\n\n".join(page_texts),
        source_kind=SourceKind.PDF,
        filename=filename,
        content_type=content_type,
        warnings=warnings,
    )


def extract_docx_text(
    file_bytes: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> ExtractedText:
    """Extract paragraphs and table text from a DOCX file."""

    if not zipfile.is_zipfile(BytesIO(file_bytes)):
        raise IngestionError("Uploaded DOCX is not a valid Office document.")

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise IngestionError("DOCX extraction requires the python-docx package.") from exc

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:  # pragma: no cover - library-specific failures
        raise IngestionError("Could not read the uploaded DOCX file.") from exc

    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return _validated_extraction(
        "\n".join(chunks),
        source_kind=SourceKind.DOCX,
        filename=filename,
        content_type=content_type,
    )


def extract_txt_text(
    file_bytes: bytes,
    *,
    filename: str | None = None,
    content_type: str | None = None,
) -> ExtractedText:
    """Decode a plain text upload."""

    decoded_text: str | None = None
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            decoded_text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue

    if decoded_text is None:
        raise IngestionError("Could not decode the uploaded text file.")

    return _validated_extraction(
        decoded_text,
        source_kind=SourceKind.TEXT,
        filename=filename,
        content_type=content_type,
    )


def validate_extracted_text(text: str, *, source_label: str = "source") -> str:
    """Normalize and validate extracted text before it enters the AI pipeline."""

    normalized = _normalize_text(text)
    if len(normalized) < MIN_EXTRACTED_TEXT_CHARS:
        raise IngestionError(
            f"Extracted text from {source_label} is too short. "
            "Please paste text manually or upload a richer document."
        )
    return normalized


def _validated_extraction(
    text: str,
    *,
    source_kind: SourceKind,
    filename: str | None = None,
    content_type: str | None = None,
    warnings: list[str] | None = None,
) -> ExtractedText:
    safe_name = safe_filename(filename, fallback="upload") if filename else None
    normalized = validate_extracted_text(text, source_label=safe_name or source_kind.value)
    return ExtractedText(
        text=normalized,
        source_kind=source_kind,
        filename=safe_name,
        content_type=content_type,
        warnings=list(warnings or []),
    )


def _normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", "").splitlines()]
    compact_lines = [line for line in lines if line]
    return "\n".join(compact_lines).strip()
